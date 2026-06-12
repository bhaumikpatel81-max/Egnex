"""
CV / Resume parser.

Tier-1 (synchronous, instant):
  - sha256 dedup hash
  - Text extraction: PDF via PyMuPDF, DOCX via python-docx (mammoth fallback),
    DOC via LibreOffice headless conversion then DOCX parse.
  - Candidate name from filename convention.
  - Skills: word-boundary match against skills_dictionary.json.
  - tsvector generation SQL fragment (done inside the INSERT, not here).

Tier-2 (async, rate-limited) is handled separately by cv_enricher.py.
"""
import hashlib
import io
import json
import os
import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Optional

# ── Skills dictionary (loaded once) ──────────────────────────────────────────

_SKILLS_PATH = Path(__file__).resolve().parent.parent / "data" / "skills_dictionary.json"
_SKILLS_RE: Optional[re.Pattern] = None

def _load_skills_re() -> re.Pattern:
    global _SKILLS_RE
    if _SKILLS_RE is None:
        with open(_SKILLS_PATH, encoding="utf-8") as f:
            data = json.load(f)
        skills = sorted(data["skills"], key=len, reverse=True)  # longer first
        # Word boundary; allow hyphens/underscores inside skill names
        parts = [re.escape(s).replace(r"\_", r"[_\-]").replace(r"\+", r"\+") for s in skills]
        pattern = r'(?<![a-zA-Z0-9_\-])(' + '|'.join(parts) + r')(?![a-zA-Z0-9_\-])'
        _SKILLS_RE = re.compile(pattern, re.IGNORECASE)
    return _SKILLS_RE


# ── Hash ──────────────────────────────────────────────────────────────────────

def sha256_hash(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


# ── Name from filename ────────────────────────────────────────────────────────

def parse_candidate_name(filename: str) -> Optional[str]:
    """
    Extract a candidate name from a filename.

    Rules (applied in order):
      1. Strip extension.
      2. Split on '_Resume' or '_resume' (case-insensitive); take the part BEFORE.
         e.g. 'John_Smith_Resume.pdf' → 'John Smith'
      3. Replace remaining underscores/hyphens with spaces, strip, title-case.
      4. If result is empty or looks like a non-name (all digits, single char), return None.
    """
    stem = Path(filename).stem  # no extension
    # Try splitting on _Resume / -Resume / _CV / -CV / _cv
    m = re.split(r'[_\-](?:resume|cv|curriculum|vitae)\b', stem, maxsplit=1, flags=re.IGNORECASE)
    name_part = m[0] if len(m) > 1 else stem

    # Replace underscores, hyphens, multiple spaces with single space
    name = re.sub(r'[_\-]+', ' ', name_part)
    name = re.sub(r'\s+', ' ', name).strip()

    # Reject if empty, too short, or all digits/non-alpha
    if not name or len(name) < 3 or not re.search(r'[A-Za-z]{2,}', name):
        return None

    return name.title()


# ── Text extraction ───────────────────────────────────────────────────────────

def _extract_pdf(data: bytes) -> str:
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=data, filetype="pdf")
        pages = [doc[i].get_text() for i in range(doc.page_count)]
        return "\n".join(pages).strip()
    except Exception:
        # Fallback to pypdf if fitz unavailable
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(data))
        return "\n".join(p.extract_text() or "" for p in reader.pages).strip()


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        parts: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = cell.text.strip()
                    if t and t not in parts:
                        parts.append(t)
        text = "\n".join(parts).strip()
        if text:
            return text
    except Exception:
        pass
    # mammoth fallback — extracts as plain text from docx
    try:
        import mammoth
        result = mammoth.extract_raw_text(io.BytesIO(data))
        return (result.value or "").strip()
    except Exception:
        pass
    return ""


def _extract_doc(data: bytes) -> str:
    """Convert .doc → .docx via LibreOffice headless, then parse as docx."""
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return ""
    with tempfile.TemporaryDirectory() as tmpdir:
        src = os.path.join(tmpdir, "input.doc")
        with open(src, "wb") as f:
            f.write(data)
        try:
            subprocess.run(
                [soffice, "--headless", "--convert-to", "docx",
                 "--outdir", tmpdir, src],
                timeout=30,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                check=True,
            )
        except Exception:
            return ""
        out_path = os.path.join(tmpdir, "input.docx")
        if not os.path.exists(out_path):
            return ""
        with open(out_path, "rb") as f:
            return _extract_docx(f.read())


def extract_text(data: bytes, ext: str) -> str:
    """
    Extract plaintext from CV bytes given the file extension (.pdf/.docx/.doc).
    Returns empty string on failure — callers decide how to handle.
    """
    ext = ext.lower().lstrip(".")
    if ext == "pdf":
        return _extract_pdf(data)
    if ext == "docx":
        return _extract_docx(data)
    if ext == "doc":
        return _extract_doc(data)
    return ""


# ── Tier-1 skill extraction ───────────────────────────────────────────────────

def extract_tier1_skills(raw_text: str) -> list[str]:
    """
    Fast keyword match against skills_dictionary.json.
    Returns sorted, deduplicated lowercase skill list.
    """
    if not raw_text:
        return []
    pattern = _load_skills_re()
    matches = pattern.findall(raw_text)
    return sorted({m.lower().replace("-", "_") for m in matches})
